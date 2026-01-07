#!/usr/bin/env python3
"""
SciENcv Biosketch Conversion Tool
---------------------------------
Converts existing NIH biosketches (Word/PDF) to SciENcv-ready format.

Key features:
1. Parses existing biosketches and extracts sections
2. Converts text to proper Markdown for SciENcv (with correct line breaks)
3. Validates character counts against new Common Form limits
4. Generates copy-paste ready text for each SciENcv field
5. Can generate Current & Pending (Other Support) XML for direct upload

Pain points addressed:
- No Word/PDF upload for biosketches (this preps text for copy-paste)
- Inconsistent line break handling (Personal Statement vs Contributions)
- Character limits validation
- Markdown formatting (**bold**, *italic*, line breaks with two spaces)
"""

import re
import json
import xml.etree.ElementTree as ET
from xml.dom import minidom
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field, asdict

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


# =============================================================================
# Character Limits for NIH Common Forms (as of January 2026)
# =============================================================================
CHAR_LIMITS = {
    "personal_statement": 3500,
    "contribution_to_science": 2000,  # per contribution, max 5
    "overall_objectives": 1500,       # for Other Support
    "honors": None,                   # max 15 entries, no char limit per entry
}


@dataclass
class BiosketchSection:
    """Represents a section of the biosketch"""
    title: str
    content: str
    citations: List[str] = field(default_factory=list)
    char_count: int = 0
    char_limit: Optional[int] = None
    warnings: List[str] = field(default_factory=list)
    
    def __post_init__(self):
        self.char_count = len(self.content)
        if self.char_limit and self.char_count > self.char_limit:
            self.warnings.append(
                f"⚠️ OVER LIMIT: {self.char_count}/{self.char_limit} characters "
                f"(reduce by {self.char_count - self.char_limit})"
            )


@dataclass
class ParsedBiosketch:
    """Complete parsed biosketch structure"""
    name: str = ""
    era_commons_id: str = ""
    position_title: str = ""
    organization: str = ""
    education: List[Dict] = field(default_factory=list)
    personal_statement: BiosketchSection = None
    positions: List[Dict] = field(default_factory=list)
    honors: List[Dict] = field(default_factory=list)
    contributions: List[BiosketchSection] = field(default_factory=list)
    research_support: List[Dict] = field(default_factory=list)
    raw_text: str = ""


# =============================================================================
# Text Extraction
# =============================================================================

def extract_text_from_docx(filepath: str) -> str:
    """Extract text from Word document"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    doc = Document(filepath)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve some structure hints
            if para.style and 'heading' in para.style.name.lower():
                paragraphs.append(f"\n### {text}\n")
            else:
                paragraphs.append(text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                paragraphs.append(" | ".join(row_text))
    
    return "\n".join(paragraphs)


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF document"""
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    
    text_parts = []
    with open(filepath, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text_parts.append(page.extract_text())
    
    return "\n".join(text_parts)


def extract_text(filepath: str) -> str:
    """Extract text from Word or PDF"""
    path = Path(filepath)
    
    if path.suffix.lower() == '.docx':
        return extract_text_from_docx(filepath)
    elif path.suffix.lower() == '.pdf':
        return extract_text_from_pdf(filepath)
    elif path.suffix.lower() == '.txt':
        return path.read_text()
    else:
        raise ValueError(f"Unsupported file type: {path.suffix}")


# =============================================================================
# Text Parsing & Section Detection
# =============================================================================

# Common section headers in NIH biosketches
SECTION_PATTERNS = {
    'personal_statement': [
        r'(?:A\.\s*)?Personal\s*Statement',
        r'Section\s*A',
    ],
    'education': [
        r'(?:B\.\s*)?Education[/\s]*Training',
        r'Professional\s*Preparation',
    ],
    'positions': [
        r'(?:B\.\s*)?Positions',
        r'Positions[,\s]*(?:Scientific\s*)?Appointments',
        r'Appointments\s*and\s*Positions',
    ],
    'honors': [
        r'Honors',
        r'Awards',
    ],
    'contributions': [
        r'(?:C\.\s*)?Contribution[s]?\s*to\s*Science',
        r'Section\s*C',
    ],
    'research_support': [
        r'(?:D\.\s*)?(?:Research|Other)\s*Support',
        r'Current\s*(?:and|&)\s*Pending',
    ],
}


def find_sections(text: str) -> Dict[str, Tuple[int, int]]:
    """Find section boundaries in text"""
    sections = {}
    
    for section_name, patterns in SECTION_PATTERNS.items():
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                sections[section_name] = match.start()
                break
    
    # Sort by position
    sorted_sections = sorted(sections.items(), key=lambda x: x[1])
    
    # Calculate end positions
    section_ranges = {}
    for i, (name, start) in enumerate(sorted_sections):
        if i + 1 < len(sorted_sections):
            end = sorted_sections[i + 1][1]
        else:
            end = len(text)
        section_ranges[name] = (start, end)
    
    return section_ranges


def extract_section_content(text: str, start: int, end: int) -> str:
    """Extract and clean section content"""
    content = text[start:end].strip()
    
    # Remove section header
    lines = content.split('\n')
    if lines:
        # Skip first line if it's the header
        first_line = lines[0].strip()
        for patterns in SECTION_PATTERNS.values():
            for pattern in patterns:
                if re.match(pattern, first_line, re.IGNORECASE):
                    lines = lines[1:]
                    break
    
    return '\n'.join(lines).strip()


def parse_contributions(content: str) -> List[BiosketchSection]:
    """Parse individual contributions to science"""
    contributions = []
    
    # Try to split by numbered contributions
    # Common patterns: "1.", "1)", "(1)", "Contribution 1"
    splits = re.split(
        r'(?:\n|^)\s*(?:(?:\d+[.\)]\s*)|(?:\(\d+\)\s*)|(?:Contribution\s*\d+[:\s]*))',
        content,
        flags=re.IGNORECASE
    )
    
    # Filter out empty splits and clean
    for i, split in enumerate(splits):
        split = split.strip()
        if split and len(split) > 50:  # Reasonable minimum for a contribution
            # Try to separate description from citations
            citations = []
            description = split
            
            # Look for citation patterns (PMIDs, DOIs, or numbered refs)
            citation_match = re.search(
                r'(?:\n\s*(?:\d+\.|[a-z]\)|\•|\-)\s*.+(?:PMID|DOI|et al|Journal|Proc|Nature|Science))',
                split,
                re.IGNORECASE
            )
            if citation_match:
                description = split[:citation_match.start()].strip()
                citations_text = split[citation_match.start():].strip()
                citations = [c.strip() for c in re.split(r'\n\s*(?:\d+\.|[a-z]\)|\•|\-)', citations_text) if c.strip()]
            
            contributions.append(BiosketchSection(
                title=f"Contribution {len(contributions) + 1}",
                content=description,
                citations=citations,
                char_limit=CHAR_LIMITS['contribution_to_science']
            ))
    
    return contributions


def parse_education_table(content: str) -> List[Dict]:
    """Parse education/training entries"""
    entries = []
    
    # Common format: INSTITUTION | DEGREE | COMPLETION DATE | FIELD
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 10:
            continue
            
        # Try to parse table-like format
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
        elif '\t' in line:
            parts = [p.strip() for p in line.split('\t')]
        else:
            # Try comma or multiple spaces
            parts = re.split(r'\s{2,}|,\s*', line)
        
        if len(parts) >= 3:
            entries.append({
                'institution': parts[0],
                'degree': parts[1] if len(parts) > 1 else '',
                'completion_date': parts[2] if len(parts) > 2 else '',
                'field': parts[3] if len(parts) > 3 else '',
            })
    
    return entries


def parse_positions(content: str) -> List[Dict]:
    """Parse positions/appointments"""
    entries = []
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 10:
            continue
        
        # Try to extract dates and position
        date_match = re.search(r'(\d{4})\s*[-–]\s*(\d{4}|present|current)', line, re.IGNORECASE)
        
        if date_match:
            dates = f"{date_match.group(1)}-{date_match.group(2)}"
            position = line[:date_match.start()].strip()
            org = line[date_match.end():].strip()
        else:
            dates = ""
            position = line
            org = ""
        
        if position:
            entries.append({
                'dates': dates,
                'position': position,
                'organization': org,
            })
    
    return entries


def parse_biosketch(text: str) -> ParsedBiosketch:
    """Parse complete biosketch text into structured format"""
    biosketch = ParsedBiosketch(raw_text=text)
    
    # Find section boundaries
    sections = find_sections(text)
    
    # Extract personal statement
    if 'personal_statement' in sections:
        start, end = sections['personal_statement']
        content = extract_section_content(text, start, end)
        biosketch.personal_statement = BiosketchSection(
            title="Personal Statement",
            content=content,
            char_limit=CHAR_LIMITS['personal_statement']
        )
    
    # Extract education
    if 'education' in sections:
        start, end = sections['education']
        content = extract_section_content(text, start, end)
        biosketch.education = parse_education_table(content)
    
    # Extract positions
    if 'positions' in sections:
        start, end = sections['positions']
        content = extract_section_content(text, start, end)
        biosketch.positions = parse_positions(content)
    
    # Extract contributions
    if 'contributions' in sections:
        start, end = sections['contributions']
        content = extract_section_content(text, start, end)
        biosketch.contributions = parse_contributions(content)
    
    return biosketch


# =============================================================================
# Markdown Conversion for SciENcv
# =============================================================================

def convert_to_sciencv_markdown(text: str, section_type: str = 'contribution') -> str:
    """
    Convert text to SciENcv-compatible Markdown format.
    
    Key differences in SciENcv Markdown:
    - Personal Statement: Enter/Return works for line breaks
    - Contributions: Must use two trailing spaces for line breaks
    - Bold: **text**
    - Italic: *text*
    - HTML tags NOT supported
    """
    result = text
    
    # Convert common formatting
    # Bold (from various Word formats)
    result = re.sub(r'\*\*\*(.+?)\*\*\*', r'**\1**', result)  # Normalize bold+italic to bold
    
    # Convert underline indicators to bold (common in Word)
    result = re.sub(r'_(.+?)_', r'*\1*', result)
    
    # Clean up multiple spaces
    result = re.sub(r' {3,}', '  ', result)
    
    # Handle line breaks based on section type
    if section_type == 'contribution':
        # Contributions need two trailing spaces for line breaks
        lines = result.split('\n')
        processed = []
        for i, line in enumerate(lines):
            line = line.rstrip()
            if line and i < len(lines) - 1:
                # Add two spaces at end for line break
                line = line + '  '
            processed.append(line)
        result = '\n'.join(processed)
    
    # Remove any HTML-like tags (not supported)
    result = re.sub(r'<[^>]+>', '', result)
    
    return result.strip()


def format_for_copy_paste(section: BiosketchSection, section_type: str = 'contribution') -> str:
    """Format a section for direct copy-paste into SciENcv"""
    output = []
    
    output.append(f"{'='*60}")
    output.append(f"SECTION: {section.title}")
    output.append(f"{'='*60}")
    
    # Show character count
    status = "✅" if not section.warnings else "⚠️"
    limit_str = f"/{section.char_limit}" if section.char_limit else ""
    output.append(f"Character count: {section.char_count}{limit_str} {status}")
    
    for warning in section.warnings:
        output.append(warning)
    
    output.append("")
    output.append("--- COPY BELOW THIS LINE ---")
    output.append("")
    
    # Converted content
    converted = convert_to_sciencv_markdown(section.content, section_type)
    output.append(converted)
    
    output.append("")
    output.append("--- END COPY ---")
    
    if section.citations:
        output.append("")
        output.append("Associated Citations (add via My Bibliography):")
        for i, cite in enumerate(section.citations, 1):
            output.append(f"  {i}. {cite}")
    
    return '\n'.join(output)


# =============================================================================
# Current & Pending (Other Support) XML Generator
# =============================================================================

def generate_cpos_xml(entries: List[Dict]) -> str:
    """
    Generate SciENcv-compatible XML for Current & Pending (Other) Support.
    
    This CAN be directly uploaded to SciENcv!
    
    Required fields per entry:
    - contributiontype: "award" (for active/pending) or "inkind"
    - title: Project title
    - status: "active" or "pending"
    - sponsor: Funding agency
    - startdate: YYYY-MM-DD
    - enddate: YYYY-MM-DD
    - personmonths: effort in person-months
    - overallobjectives: Brief description (max 1500 chars)
    """
    
    # XML namespaces
    nsmap = {
        'xmlns': 'http://www.ncbi.nlm.nih.gov/sciencv',
    }
    
    root = ET.Element('sciencv', nsmap)
    
    # Add identifying info section (optional but helpful)
    id_info = ET.SubElement(root, 'identifying-information')
    
    # Add support entries
    support_section = ET.SubElement(root, 'current-and-pending-support')
    
    for entry in entries:
        support = ET.SubElement(support_section, 'support')
        
        # Required: contribution type
        ctype = ET.SubElement(support, 'contributiontype')
        ctype.text = entry.get('contributiontype', 'award')
        
        # Title
        title = ET.SubElement(support, 'title')
        title.text = entry.get('title', '')
        
        # Status (active/pending)
        status = ET.SubElement(support, 'status')
        status.text = entry.get('status', 'active')
        
        # Sponsor
        sponsor = ET.SubElement(support, 'sponsor')
        sponsor.text = entry.get('sponsor', '')
        
        # Award number
        if entry.get('award_number'):
            award_num = ET.SubElement(support, 'awardnumber')
            award_num.text = entry['award_number']
        
        # Dates
        if entry.get('startdate'):
            start = ET.SubElement(support, 'startdate')
            start.text = entry['startdate']
        
        if entry.get('enddate'):
            end = ET.SubElement(support, 'enddate')
            end.text = entry['enddate']
        
        # Person months
        if entry.get('personmonths'):
            pm = ET.SubElement(support, 'personmonths')
            # Can be broken down by year
            year = ET.SubElement(pm, 'year')
            year.set('value', str(datetime.now().year))
            year.text = str(entry['personmonths'])
        
        # Overall objectives (max 1500 chars)
        if entry.get('overallobjectives'):
            obj = ET.SubElement(support, 'overallobjectives')
            obj_text = entry['overallobjectives'][:1500]  # Enforce limit
            obj.text = obj_text
        
        # PI info
        if entry.get('pi_name'):
            pi = ET.SubElement(support, 'pi')
            pi.text = entry['pi_name']
    
    # Pretty print
    xml_str = ET.tostring(root, encoding='unicode')
    dom = minidom.parseString(xml_str)
    return dom.toprettyxml(indent="  ")


# =============================================================================
# Main Conversion Functions
# =============================================================================

def convert_biosketch(input_file: str, output_dir: str = None) -> Dict:
    """
    Main conversion function.
    
    Args:
        input_file: Path to Word (.docx), PDF, or text file
        output_dir: Output directory (defaults to same as input)
    
    Returns:
        Dict with conversion results and output paths
    """
    input_path = Path(input_file)
    
    if output_dir:
        out_path = Path(output_dir)
    else:
        out_path = input_path.parent
    
    out_path.mkdir(parents=True, exist_ok=True)
    
    # Extract text
    print(f"📄 Extracting text from: {input_path.name}")
    raw_text = extract_text(str(input_path))
    
    # Parse structure
    print("🔍 Parsing biosketch structure...")
    biosketch = parse_biosketch(raw_text)
    
    # Generate outputs
    results = {
        'input_file': str(input_path),
        'outputs': [],
        'warnings': [],
        'sections_found': [],
    }
    
    # Main output file with all sections
    main_output = []
    main_output.append("=" * 70)
    main_output.append("SciENcv BIOSKETCH CONVERSION RESULTS")
    main_output.append(f"Source: {input_path.name}")
    main_output.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    main_output.append("=" * 70)
    main_output.append("")
    main_output.append("INSTRUCTIONS:")
    main_output.append("1. Log into SciENcv via My NCBI")
    main_output.append("2. Create a new NIH Biographical Sketch Common Form")
    main_output.append("3. For each section below, copy the text between the markers")
    main_output.append("4. Paste into the corresponding SciENcv field")
    main_output.append("")
    main_output.append("IMPORTANT MARKDOWN NOTES:")
    main_output.append("- Personal Statement: Press Enter for line breaks (normal)")
    main_output.append("- Contributions: Line breaks need TWO TRAILING SPACES")
    main_output.append("  (this conversion adds them automatically)")
    main_output.append("- Bold: **text**  Italic: *text*")
    main_output.append("- HTML tags are NOT supported")
    main_output.append("")
    
    # Personal Statement
    if biosketch.personal_statement:
        results['sections_found'].append('personal_statement')
        main_output.append(format_for_copy_paste(
            biosketch.personal_statement, 
            section_type='personal_statement'
        ))
        main_output.append("\n")
        
        if biosketch.personal_statement.warnings:
            results['warnings'].extend(biosketch.personal_statement.warnings)
    
    # Contributions to Science
    for i, contrib in enumerate(biosketch.contributions):
        results['sections_found'].append(f'contribution_{i+1}')
        main_output.append(format_for_copy_paste(contrib, section_type='contribution'))
        main_output.append("\n")
        
        if contrib.warnings:
            results['warnings'].extend(contrib.warnings)
    
    # Education (as reference)
    if biosketch.education:
        results['sections_found'].append('education')
        main_output.append("=" * 60)
        main_output.append("EDUCATION/TRAINING (enter via SciENcv form)")
        main_output.append("=" * 60)
        for edu in biosketch.education:
            main_output.append(f"  Institution: {edu.get('institution', '')}")
            main_output.append(f"  Degree: {edu.get('degree', '')}")
            main_output.append(f"  Date: {edu.get('completion_date', '')}")
            main_output.append(f"  Field: {edu.get('field', '')}")
            main_output.append("")
        main_output.append("\n")
    
    # Positions (as reference)
    if biosketch.positions:
        results['sections_found'].append('positions')
        main_output.append("=" * 60)
        main_output.append("POSITIONS/APPOINTMENTS (enter via SciENcv form)")
        main_output.append("=" * 60)
        for pos in biosketch.positions:
            main_output.append(f"  {pos.get('dates', '')} | {pos.get('position', '')} | {pos.get('organization', '')}")
        main_output.append("\n")
    
    # Write main output
    main_output_file = out_path / f"{input_path.stem}_sciencv_ready.txt"
    main_output_file.write_text('\n'.join(main_output))
    results['outputs'].append(str(main_output_file))
    
    print(f"✅ Main output: {main_output_file.name}")
    
    # Summary
    print("\n" + "=" * 50)
    print("CONVERSION SUMMARY")
    print("=" * 50)
    print(f"Sections found: {', '.join(results['sections_found'])}")
    
    if results['warnings']:
        print("\n⚠️  WARNINGS:")
        for w in results['warnings']:
            print(f"  {w}")
    else:
        print("\n✅ All sections within character limits")
    
    return results


def convert_other_support_to_xml(entries: List[Dict], output_file: str) -> str:
    """
    Convert Other Support entries to uploadable XML.
    
    Example entry format:
    {
        'contributiontype': 'award',  # or 'inkind'
        'title': 'Project Title',
        'status': 'active',  # or 'pending'
        'sponsor': 'NIH/NHLBI',
        'award_number': 'R01 HL123456',
        'startdate': '2023-09-01',
        'enddate': '2028-08-31',
        'personmonths': 3.6,
        'overallobjectives': 'Brief description of project goals...',
        'pi_name': 'John Smith, PhD'
    }
    """
    xml_content = generate_cpos_xml(entries)
    
    out_path = Path(output_file)
    out_path.write_text(xml_content)
    
    print(f"✅ Generated CPOS XML: {out_path}")
    print(f"   Upload this file directly to SciENcv!")
    
    return xml_content


# =============================================================================
# CLI Interface
# =============================================================================

def main():
    import sys
    
    print("""
╔════════════════════════════════════════════════════════════════╗
║           SciENcv Biosketch Conversion Tool                    ║
║                                                                ║
║  Converts Word/PDF biosketches to SciENcv-ready format         ║
║  Handles Markdown formatting and character limits              ║
╚════════════════════════════════════════════════════════════════╝
    """)
    
    if len(sys.argv) < 2:
        print("Usage: python sciencv_converter.py <biosketch_file> [output_dir]")
        print("")
        print("Supported formats: .docx, .pdf, .txt")
        print("")
        print("Example:")
        print("  python sciencv_converter.py my_biosketch.docx")
        print("  python sciencv_converter.py my_biosketch.pdf ./output/")
        return
    
    input_file = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        results = convert_biosketch(input_file, output_dir)
        print("\n🎉 Conversion complete!")
        print(f"Output file: {results['outputs'][0]}")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
